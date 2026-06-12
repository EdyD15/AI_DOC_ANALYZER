import os, base64, logging, io, json, sqlite3
from datetime import datetime, timezone

import bcrypt

def _get_api_key() -> str:
    key = os.getenv("OPENAI_API_KEY")
    if key:
        return key
    try:
        import streamlit as st
        return st.secrets["OPENAI_API_KEY"]
    except Exception:
        raise RuntimeError("OPENAI_API_KEY not set. Add it to .streamlit/secrets.toml or set the env var.")

from pypdf import PdfReader
import docx
import chromadb
from langchain_text_splitters import RecursiveCharacterTextSplitter
from chromadb.utils import embedding_functions
from openai import OpenAI
from fpdf import FPDF
from langchain_openai import ChatOpenAI
from langchain_core.tools import tool
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage, ToolMessage

# ==========================================
# CONFIGURATION
# ==========================================

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
logger = logging.getLogger(__name__)

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
VECTOR_DB_PATH = os.path.join(_BASE_DIR, "vector_db")
CHAT_DB_PATH = os.path.join(_BASE_DIR, "chat_history.db")

CHUNK_SIZE = 1000
OVERLAP = 200
N_RESULTS = 4
MAX_TOOL_ITERATIONS = 5

# ==========================================
# 1. VECTOR DATABASE CONNECTION
# ==========================================

def get_chroma_collection(user_id):
    api_key = _get_api_key()
    openai_ef = embedding_functions.OpenAIEmbeddingFunction(
        api_key=api_key, model_name="text-embedding-3-small"
    )
    client = chromadb.PersistentClient(path=VECTOR_DB_PATH)
    collection = client.get_or_create_collection(
        name=f"corporate_docs_{user_id}", embedding_function=openai_ef
    )
    return collection

def chunk_text(text, chunk_size=CHUNK_SIZE, overlap=OVERLAP):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", "! ", "? ", " ", ""],
    )
    return splitter.split_text(text)

def clear_db(user_id):
    try:
        client = chromadb.PersistentClient(path=VECTOR_DB_PATH)
        client.delete_collection(f"corporate_docs_{user_id}")
    except Exception as e:
        logger.warning("Could not clear vector DB: %s", e)

def get_all_documents_from_db(user_id):
    try:
        collection = get_chroma_collection(user_id)
        result = collection.get(include=["metadatas"])
        unique_files = {meta["source"] for meta in result["metadatas"] if "source" in meta}
        return {file: "Vectorized" for file in unique_files}
    except Exception as e:
        logger.error("Failed to fetch documents from ChromaDB: %s", e)
        return {}

# ==========================================
# 2. READING, CHUNKING, AND PAGE TRACKING
# ==========================================

_IMAGE_MIME = {
    "jpg": "image/jpeg",
    "jpeg": "image/jpeg",
    "png": "image/png",
    "webp": "image/webp",
}
_IMAGE_MAX_BYTES = 20 * 1024 * 1024  # 20 MB — OpenAI base64 limit

def _decode_text_file(file_bytes):
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return file_bytes.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return file_bytes.decode("latin-1", errors="replace")

def _extract_text_from_image(uploaded_file):
    """Send an image to GPT-4o Vision and return a text description of all content."""
    raw = uploaded_file.read()
    if len(raw) > _IMAGE_MAX_BYTES:
        raise ValueError(f"Image exceeds 20 MB limit ({len(raw) // (1024*1024)} MB).")

    ext = uploaded_file.name.lower().rsplit(".", 1)[-1]
    mime_type = _IMAGE_MIME.get(ext, "image/jpeg")
    b64 = base64.b64encode(raw).decode("utf-8")

    api_key = _get_api_key()
    client = OpenAI(api_key=api_key)

    response = client.chat.completions.create(
        model="gpt-4o",
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "You are a document analysis assistant. "
                        "First, extract ALL text visible in this image exactly as it appears. "
                        "Then describe any charts, diagrams, tables, or other visual elements in detail. "
                        "Be thorough — this output will be used to answer questions about the image."
                    )
                },
                {
                    "type": "image_url",
                    "image_url": {"url": f"data:{mime_type};base64,{b64}"}
                }
            ]
        }],
        max_tokens=4096
    )
    return response.choices[0].message.content

def process_and_save_document(uploaded_file, user_id):
    """Returns (ok, error_message). error_message is None on success."""
    name = uploaded_file.name
    file_name = name.lower()

    try:
        collection = get_chroma_collection(user_id)
    except Exception as e:
        logger.error("Could not connect to vector database: %s", e)
        return False, "Could not connect to the vector database."

    existing = collection.get(where={"source": name})
    if existing["ids"]:
        return True, None

    pages_data = []
    looks_scanned = False

    try:
        if file_name.endswith(".pdf"):
            reader = PdfReader(uploaded_file)
            for i, page in enumerate(reader.pages):
                ext = page.extract_text()
                if ext and ext.strip():
                    pages_data.append({"text": ext, "page": i + 1})
                elif page.images:
                    looks_scanned = True

        elif file_name.endswith(".docx"):
            doc = docx.Document(uploaded_file)
            full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            if full_text:
                pages_data.append({"text": full_text, "page": 1})

        elif file_name.endswith(".txt"):
            raw = uploaded_file.read()
            text = _decode_text_file(raw)
            if text:
                pages_data.append({"text": text, "page": 1})

        elif file_name.endswith((".png", ".jpg", ".jpeg", ".webp")):
            text = _extract_text_from_image(uploaded_file)
            if text:
                pages_data.append({"text": text, "page": 1})

    except Exception as e:
        logger.error("File read error for %s: %s", name, e)
        return False, "Could not read the file."

    if not pages_data:
        if looks_scanned:
            logger.warning("%s appears to be a scanned PDF with no text layer.", name)
            return False, (
                "This PDF appears to be scanned (image-only pages, no extractable text). "
                "Convert its pages to images (PNG/JPG) and upload them individually instead "
                "— those go through OCR automatically."
            )
        logger.warning("No text could be extracted from %s.", name)
        return False, "No text could be extracted from this file."

    all_chunks, all_metadatas, all_ids = [], [], []
    chunk_id_counter = 0

    for p_data in pages_data:
        for c in chunk_text(p_data["text"]):
            all_chunks.append(c)
            all_metadatas.append({"source": name, "page": p_data["page"]})
            all_ids.append(f"{name}_chunk_{chunk_id_counter}")
            chunk_id_counter += 1

    if all_chunks:
        try:
            collection.add(documents=all_chunks, metadatas=all_metadatas, ids=all_ids)
        except Exception as e:
            logger.error("ChromaDB add error for %s: %s", name, e)
            return False, "Could not save the document to the vector database."

    return True, None

# ==========================================
# 3. QUERYING AND CITATIONS (LANGCHAIN AGENT)
# ==========================================

AGENT_SYSTEM_PROMPT = (
    "You are an advanced corporate document analyst with access to tools for working with "
    "the user's knowledge base.\n"
    "- Use 'search_knowledge_base' to find relevant excerpts before answering any question "
    "that could be answered from the user's documents.\n"
    "- Use 'list_documents' when the user asks what documents are available.\n"
    "- Use 'summarize_document' when the user asks for a summary of a specific document "
    "(pass its exact file name, as returned by 'list_documents').\n\n"
    "NOTE: Uploaded images are stored as text descriptions (extracted via OCR/vision when the "
    "image was uploaded). If the user asks what an image shows, treat the matching excerpt's "
    "text as your description of that image's visual content — do not claim you cannot view "
    "images.\n\n"
    "CRITICAL RULE: Whenever you use information returned by 'search_knowledge_base' or "
    "'summarize_document', you MUST explicitly cite the exact Source and Page shown in the "
    "excerpt headers. Format citations at the end of the relevant sentence like this: "
    "'\U0001f4dd [Document: X, Page: Y]'. "
    "If the knowledge base has no relevant information, do not hallucinate. Start your "
    "answer with '\U0001f310 [General Knowledge]:' and answer using your general knowledge.\n\n"
    "LANGUAGE RULE: Always reply in the same language the user's question is written in, "
    "regardless of the language of the source documents. Translate any cited content into "
    "the user's language while preserving the citation.\n\n"
    "MULTI-DOCUMENT RULE: When an answer draws on multiple documents, never blend their "
    "information into a single undifferentiated statement. Address each document's "
    "contribution separately (e.g. in its own sentence or bullet point), each with its own "
    "'\U0001f4dd [Document: X, Page: Y]' citation, so the user can tell which document each "
    "piece of information came from."
)


def _build_tools(user_id, selected_documents=None):
    def _where_filter():
        if not selected_documents:
            return None
        if len(selected_documents) == 1:
            return {"source": selected_documents[0]}
        return {"source": {"$in": selected_documents}}

    @tool
    def search_knowledge_base(query: str) -> str:
        """Search the document knowledge base for excerpts relevant to the query.
        Returns excerpts prefixed with their source file name and page number.
        Use this for any question that might be answered from the user's uploaded documents."""
        try:
            collection = get_chroma_collection(user_id)
            results = collection.query(
                query_texts=[query],
                n_results=N_RESULTS,
                where=_where_filter(),
                include=["documents", "metadatas"],
            )
        except Exception as e:
            logger.error("ChromaDB query failed: %s", e)
            return "Document search failed."

        if not results["documents"] or not results["documents"][0]:
            return "No relevant excerpts found."

        parts = []
        for doc, meta in zip(results["documents"][0], results["metadatas"][0]):
            source = meta.get("source", "Unknown")
            page = meta.get("page", "N/A")
            parts.append(f"--- [SOURCE: {source} | PAGE: {page}] ---\n{doc}")
        return "\n\n".join(parts)

    @tool
    def list_documents() -> str:
        """List all documents currently stored in the knowledge base."""
        docs = get_all_documents_from_db(user_id)
        if not docs:
            return "The knowledge base is empty."
        return ", ".join(sorted(docs.keys()))

    @tool
    def summarize_document(document_name: str) -> str:
        """Retrieve the full text of a specific document so it can be summarized.
        Pass the exact file name as returned by 'list_documents'."""
        try:
            collection = get_chroma_collection(user_id)
            result = collection.get(where={"source": document_name}, include=["documents", "metadatas"])
        except Exception as e:
            logger.error("ChromaDB get failed: %s", e)
            return "Could not retrieve the document."

        if not result["ids"]:
            return f"No document named '{document_name}' found in the knowledge base."

        pairs = sorted(zip(result["documents"], result["metadatas"]), key=lambda x: x[1].get("page", 0))
        text = "\n\n".join(f"[Page {m.get('page', 'N/A')}]\n{d}" for d, m in pairs)

        if len(text) > 12000:
            text = text[:12000] + "\n...[content truncated]"
        return f"--- [SOURCE: {document_name}] ---\n{text}"

    return [search_knowledge_base, list_documents, summarize_document]


def _history_to_messages(chat_history):
    messages = []
    for m in chat_history or []:
        role = m.get("role")
        content = m.get("content", "")
        if not content:
            continue
        if role == "user":
            messages.append(HumanMessage(content=content))
        elif role == "assistant":
            messages.append(AIMessage(content=content))
    return messages


def query_openai_api(user_question, user_id, selected_documents=None, image_base64=None, image_mime=None, chat_history=None):
    api_key = _get_api_key()

    has_image = bool(image_base64 and image_mime)
    model = "gpt-4o" if has_image else "gpt-4o-mini"

    llm = ChatOpenAI(model=model, api_key=api_key, streaming=True)
    tools = _build_tools(user_id, selected_documents)
    llm_with_tools = llm.bind_tools(tools)
    llm_force_search = llm.bind_tools(
        tools, tool_choice={"type": "function", "function": {"name": "search_knowledge_base"}}
    )
    llm_force_any = llm.bind_tools(tools, tool_choice="required")
    tool_map = {t.name: t for t in tools}

    system_prompt = AGENT_SYSTEM_PROMPT
    if selected_documents:
        system_prompt += (
            "\n\nThe user has scoped this conversation to the following document(s): "
            f"{', '.join(selected_documents)}. 'search_knowledge_base' is automatically "
            "restricted to these documents."
        )

    messages = [SystemMessage(content=system_prompt)]
    messages.extend(_history_to_messages(chat_history))

    if has_image:
        user_content = [
            {"type": "text", "text": user_question},
            {"type": "image_url", "image_url": {"url": f"data:{image_mime};base64,{image_base64}"}},
        ]
    else:
        user_content = user_question
    messages.append(HumanMessage(content=user_content))

    try:
        for i in range(MAX_TOOL_ITERATIONS):
            if i == 0:
                # When the conversation is scoped to specific document(s) or an image was
                # just attached, the user is almost certainly asking about that content —
                # force a knowledge-base search so it isn't skipped in favor of another tool.
                step_llm = llm_force_search if (has_image or selected_documents) else llm_force_any
            else:
                step_llm = llm_with_tools
            full_chunk = None
            for chunk in step_llm.stream(messages):
                if chunk.content:
                    yield chunk.content
                full_chunk = chunk if full_chunk is None else full_chunk + chunk

            if full_chunk is None:
                break

            messages.append(full_chunk)

            if not full_chunk.tool_calls:
                break

            for tool_call in full_chunk.tool_calls:
                tool_fn = tool_map.get(tool_call["name"])
                if tool_fn:
                    try:
                        result = tool_fn.invoke(tool_call["args"])
                    except Exception as e:
                        result = f"Tool error: {e}"
                else:
                    result = f"Unknown tool: {tool_call['name']}"
                messages.append(ToolMessage(content=str(result), tool_call_id=tool_call["id"]))
        else:
            yield "\n\n⚠️ Reached the maximum number of tool calls without a final answer."
    except Exception as e:
        logger.error("OpenAI API error: %s", e)
        raise RuntimeError(f"AI service error: {e}") from e

# ==========================================
# 4. CHAT EXPORT TOOL
# ==========================================

def generate_chat_export_docx(messages):
    doc = docx.Document()
    doc.add_heading("DocuMind AI - Session Export", 0)

    for msg in messages:
        role = "\U0001f464 You" if msg["role"] == "user" else "\U0001f916 DocuMind AI"
        p = doc.add_paragraph()
        p.add_run(f"{role}:\n").bold = True
        p.add_run(msg["content"])
        doc.add_paragraph("-" * 40)

    byte_io = io.BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

def generate_chat_export_pdf(messages):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)

    pdf.set_font("Arial", "B", 16)
    pdf.cell(0, 10, "DocuMind AI - Session Export", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)

    for msg in messages:
        role = "You" if msg["role"] == "user" else "DocuMind AI"

        pdf.set_font("Arial", "B", 12)
        pdf.cell(0, 8, f"{role}:", new_x="LMARGIN", new_y="NEXT")

        pdf.set_font("Arial", "", 12)
        safe_text = msg["content"].encode("latin-1", "replace").decode("latin-1")
        pdf.multi_cell(0, 8, safe_text)
        pdf.ln(5)

    byte_io = io.BytesIO()
    byte_io.write(pdf.output())
    byte_io.seek(0)
    return byte_io

# ==========================================
# 5. PERMANENT CHAT HISTORY (SQLite)
# ==========================================

def init_chat_db():
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        conn.execute(
            "CREATE TABLE IF NOT EXISTS users ("
            "id INTEGER PRIMARY KEY AUTOINCREMENT, "
            "username TEXT UNIQUE NOT NULL, "
            "password_hash TEXT NOT NULL, "
            "created_at TEXT NOT NULL)"
        )

        conn.execute(
            "CREATE TABLE IF NOT EXISTS chat_sessions ("
            "user_id INTEGER NOT NULL, "
            "session_name TEXT NOT NULL, "
            "messages TEXT, "
            "PRIMARY KEY (user_id, session_name))"
        )

        # Migrate the legacy single-user schema (session_name TEXT PRIMARY KEY,
        # no user_id) — pre-auth sessions weren't tied to any account, so they
        # are dropped rather than migrated.
        cols = [row[1] for row in conn.execute("PRAGMA table_info(chat_sessions)").fetchall()]
        if "user_id" not in cols:
            conn.execute("DROP TABLE chat_sessions")
            conn.execute(
                "CREATE TABLE chat_sessions ("
                "user_id INTEGER NOT NULL, "
                "session_name TEXT NOT NULL, "
                "messages TEXT, "
                "PRIMARY KEY (user_id, session_name))"
            )

        conn.commit()

def save_chat_session(user_id, session_name, messages):
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        conn.execute(
            "INSERT OR REPLACE INTO chat_sessions (user_id, session_name, messages) VALUES (?, ?, ?)",
            (user_id, session_name, json.dumps(messages))
        )
        conn.commit()

def load_all_chat_sessions(user_id):
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        rows = conn.execute(
            "SELECT session_name, messages FROM chat_sessions WHERE user_id = ?",
            (user_id,)
        ).fetchall()
    return {row[0]: json.loads(row[1]) for row in rows}

def delete_chat_session(user_id, session_name):
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        conn.execute(
            "DELETE FROM chat_sessions WHERE user_id = ? AND session_name = ?",
            (user_id, session_name)
        )
        conn.commit()

# ==========================================
# 6. USERS (SQLite)
# ==========================================

def create_user(username, password):
    password_hash = bcrypt.hashpw(password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
    try:
        with sqlite3.connect(CHAT_DB_PATH) as conn:
            cur = conn.execute(
                "INSERT INTO users (username, password_hash, created_at) VALUES (?, ?, ?)",
                (username, password_hash, datetime.now(timezone.utc).isoformat())
            )
            conn.commit()
            return {"id": cur.lastrowid, "username": username}
    except sqlite3.IntegrityError:
        return None

def authenticate_user(username, password):
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        row = conn.execute(
            "SELECT id, username, password_hash FROM users WHERE username = ?",
            (username,)
        ).fetchone()

    if row is None:
        return None

    user_id, db_username, password_hash = row
    if not bcrypt.checkpw(password.encode("utf-8"), password_hash.encode("utf-8")):
        return None
    return {"id": user_id, "username": db_username}

def change_password(user_id, current_password, new_password):
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        row = conn.execute("SELECT password_hash FROM users WHERE id = ?", (user_id,)).fetchone()
        if row is None:
            return False

        if not bcrypt.checkpw(current_password.encode("utf-8"), row[0].encode("utf-8")):
            return False

        new_hash = bcrypt.hashpw(new_password.encode("utf-8"), bcrypt.gensalt()).decode("utf-8")
        conn.execute("UPDATE users SET password_hash = ? WHERE id = ?", (new_hash, user_id))
        conn.commit()
    return True

def get_user_by_id(user_id):
    with sqlite3.connect(CHAT_DB_PATH) as conn:
        row = conn.execute("SELECT id, username FROM users WHERE id = ?", (user_id,)).fetchone()
    if row is None:
        return None
    return {"id": row[0], "username": row[1]}

# Initialize the chat database when the module loads
init_chat_db()
